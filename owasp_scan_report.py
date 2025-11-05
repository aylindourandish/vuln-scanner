#!/usr/bin/env python3
#cd C:\Users\dorandish\vuln_scanner
#dir
#python .\vulnweb_scanner.py
#https://example.test/
#http://localhost:8000/

#!/usr/bin/env python3
import os
import sys
import json
import datetime
import re
import shlex
import subprocess
from urllib.parse import urljoin, urlparse, parse_qs
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

requests.packages.urllib3.disable_warnings(category=Warning)

TIMEOUT = 15
USER_AGENT = "Mozilla/5.0 (X11; Linux x86_64) owasp-scan/1.0"
HEADERS = {"User-Agent": USER_AGENT}
RESULTS_DIR = "results_owasp"

def run_cmd(cmd, timeout=600):
    try:
        print("[*] CMD:", cmd)
        p = subprocess.run(shlex.split(cmd), stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, timeout=timeout)
        return p.returncode, p.stdout
    except Exception as e:
        return 1, str(e)

def ensure_dir(p):
    os.makedirs(p, exist_ok=True)

def fetch_url(url, method="GET", data=None, headers=None, allow_redirects=True):
    try:
        h = headers or HEADERS
        if method.upper() == "GET":
            r = requests.get(url, headers=h, timeout=TIMEOUT, verify=False, allow_redirects=allow_redirects)
        else:
            r = requests.post(url, data=data, headers=h, timeout=TIMEOUT, verify=False, allow_redirects=allow_redirects)
        return r
    except Exception:
        return None

def basic_tech_info(target, outdir):
    summary = {}
    parsed = urlparse(target)
    host = parsed.hostname or target
    rc, out = run_cmd(f"whatweb -v {target}")
    try:
        with open(os.path.join(outdir, "whatweb.txt"), "w", encoding="utf-8") as fh:
            fh.write(out)
    except Exception:
        pass
    summary['whatweb'] = out.splitlines()[:10] if out else []
    rc, out = run_cmd(f"nmap -sV -p 80,443 -Pn {host} -oN {os.path.join(outdir,'nmap_services.txt')}", timeout=300)
    summary['nmap'] = out.splitlines()[:30] if out else []
    try:
        with open(os.path.join(outdir, "nmap_services.txt"), "w", encoding="utf-8") as fh:
            fh.write(out or "")
    except Exception:
        pass
    return summary

def check_headers_tls(target, outdir):
    findings = {}
    r = fetch_url(target)
    if not r:
        findings['headers_error'] = "No response"
        return findings
    headers = dict(r.headers)
    try:
        with open(os.path.join(outdir, "http_headers_raw.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(headers, indent=2, ensure_ascii=False))
    except Exception:
        pass
    findings['headers'] = headers
    parsed = urlparse(target)
    hostname = parsed.hostname
    if parsed.scheme == "https" or (hostname and hostname.endswith(":443")) or ":443" in target:
        rc, out = run_cmd(f"nmap -p 443 --script ssl-enum-ciphers {hostname}", timeout=200)
        try:
            with open(os.path.join(outdir, "nmap_ssl.txt"), "w", encoding="utf-8") as fh:
                fh.write(out or "")
        except Exception:
            pass
        findings['ssl_enum_head'] = out.splitlines()[:40] if out else []
        try:
            cert = requests.get(target, headers=HEADERS, timeout=TIMEOUT, verify=True)
            findings['tls_verify'] = "cert verified"
        except requests.exceptions.SSLError as e:
            findings['tls_verify'] = f"SSL verify error: {e}"
        except Exception as e:
            findings['tls_verify'] = f"TLS check error: {e}"
    return findings

def check_sensitive_files(target, outdir):
    sensitive = ["/.git/", "/.git/config", "/.env", "/config.php", "/.htpasswd", "/backup.zip"]
    found = {}
    for p in sensitive:
        url = urljoin(target, p)
        r = fetch_url(url)
        if r and r.status_code in (200, 403, 301, 302):
            if r.status_code == 200 and len(r.text or "") > 50:
                found[p] = {'status': r.status_code, 'snippet': (r.text or "")[:300]}
            else:
                found[p] = {'status': r.status_code, 'note': 'present or redirected'}
    try:
        with open(os.path.join(outdir, "sensitive_files.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(found, indent=2, ensure_ascii=False))
    except Exception:
        pass
    return found

def find_links(target, outdir):
    r = fetch_url(target)
    links = set()
    forms = []
    if not r:
        try:
            with open(os.path.join(outdir, "links_and_forms.json"), "w", encoding="utf-8") as fh:
                fh.write(json.dumps({'links': [], 'forms': []}, indent=2, ensure_ascii=False))
        except Exception:
            pass
        return links, forms
    soup = BeautifulSoup(r.text or "", "lxml")
    for a in soup.find_all("a", href=True):
        href = a['href']
        full = urljoin(target, href)
        links.add(full)
    for f in soup.find_all("form"):
        action = f.get('action') or target
        action = urljoin(target, action)
        method = f.get('method', 'get').lower()
        inputs = []
        for inp in f.find_all(["input", "textarea", "select"]):
            name = inp.get('name')
            if not name:
                continue
            inputs.append({'name': name, 'type': inp.get('type', 'text')})
        forms.append({'action': action, 'method': method, 'inputs': inputs})
    try:
        with open(os.path.join(outdir, "links_and_forms.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps({'links': list(links), 'forms': forms}, indent=2, ensure_ascii=False))
    except Exception:
        pass
    return links, forms

SQL_ERRORS = [
    "you have an error in your sql syntax",
    "warning: mysql",
    "unclosed quotation mark after the character string",
    "quoted string not properly terminated",
    "pg_query():",
    "mysql_fetch",
    "syntax error"
]

def check_sqli_basic(target, links, outdir):
    findings = []
    payloads = ["' OR '1'='1", "\" OR \"1\"=\"1", "' OR '1'='1' -- -"]
    for url in list(links)[:80]:
        parsed = urlparse(url)
        qs = parse_qs(parsed.query)
        if not qs:
            continue
        for param in qs:
            original = qs[param][0] if qs[param] else ""
            for p in payloads:
                new_qs = parsed.query.replace(original, original + p)
                test_url = parsed._replace(query=new_qs).geturl()
                r = fetch_url(test_url)
                if not r:
                    continue
                text = (r.text or "").lower()
                for sig in SQL_ERRORS:
                    if sig in text:
                        findings.append({'url': test_url, 'evidence_snippet': (r.text or "")[:500], 'signature': sig})
    try:
        with open(os.path.join(outdir, "sqli_findings.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(findings, indent=2, ensure_ascii=False))
    except Exception:
        pass
    return findings

XSS_PROBES = ["<script>alert('xss1')</script>", "\"><script>alert(1)</script>"]

def check_xss_reflected(target, links, forms, outdir):
    findings = []
    for url in list(links)[:80]:
        parsed = urlparse(url)
        qs = parse_qs(parsed.query)
        if not qs:
            continue
        for param in qs:
            for p in XSS_PROBES:
                new_qs = parsed.query.replace(qs[param][0], qs[param][0] + p)
                test_url = parsed._replace(query=new_qs).geturl()
                r = fetch_url(test_url)
                if not r:
                    continue
                if p in (r.text or ""):
                    findings.append({'type': 'reflected', 'url': test_url, 'payload': p, 'status': r.status_code})
    for f in forms:
        for inp in f['inputs'][:6]:
            name = inp['name']
            for p in XSS_PROBES:
                data = {name: p}
                r = fetch_url(f['action'], method=f['method'].upper(), data=data)
                if not r:
                    continue
                if p in (r.text or ""):
                    findings.append({'type': 'form', 'action': f['action'], 'input': name, 'payload': p})
    try:
        with open(os.path.join(outdir, "xss_findings.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(findings, indent=2, ensure_ascii=False))
    except Exception:
        pass
    return findings

def check_idor(target, links, outdir):
    findings = []
    for url in list(links)[:120]:
        parsed = urlparse(url)
        qs = parse_qs(parsed.query)
        if not qs:
            continue
        for param in qs:
            if param.lower() in ('id', 'user', 'uid', 'profile', 'account', 'page'):
                val = qs[param][0]
                if val.isdigit():
                    base_r = fetch_url(url)
                    base_len = len(base_r.text or "") if base_r else 0
                    for delta in (1, 2, -1):
                        new_val = str(int(val) + delta)
                        new_q = parsed.query.replace(val, new_val)
                        test_url = parsed._replace(query=new_q).geturl()
                        r = fetch_url(test_url)
                        if not r:
                            continue
                        if abs(len(r.text or "") - base_len) > 50:
                            findings.append({'url': test_url, 'param': param, 'new_val': new_val, 'len': len(r.text or "")})
    try:
        with open(os.path.join(outdir, "idor_findings.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(findings, indent=2, ensure_ascii=False))
    except Exception:
        pass
    return findings

def check_csrf(forms, outdir):
    missing = []
    for f in forms:
        inputs = [i['name'] for i in f['inputs']]
        if not any('csrf' in (n or '').lower() or 'token' in (n or '').lower() for n in inputs):
            missing.append({'action': f['action'], 'method': f['method'], 'inputs': inputs})
    try:
        with open(os.path.join(outdir, "csrf_findings.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(missing, indent=2, ensure_ascii=False))
    except Exception:
        pass
    return missing

def check_security_misconfig(headers, outdir):
    issues = []
    lower_keys = {k.lower() for k in headers}
    if 'strict-transport-security' not in lower_keys:
        issues.append('Missing HSTS')
    if 'x-frame-options' not in lower_keys:
        issues.append('Missing X-Frame-Options')
    if 'content-security-policy' not in lower_keys:
        issues.append('Missing Content-Security-Policy (CSP)')
    cookies = headers.get('set-cookie', '')
    issues.append({'set-cookie': cookies})
    try:
        with open(os.path.join(outdir, "misconfig_findings.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(issues, indent=2, ensure_ascii=False))
    except Exception:
        pass
    return issues

def shutil_which(name):
    from shutil import which
    return which(name) is not None

def aggressive_checks(target, outdir):
    ag = {}
    hostname = urlparse(target).hostname
    rc, out = run_cmd(f"nmap --script vuln -p 80,443 {hostname}", timeout=600)
    try:
        with open(os.path.join(outdir, "nmap_vuln.txt"), "w", encoding="utf-8") as fh:
            fh.write(out or "")
    except Exception:
        pass
    ag['nmap_vuln_head'] = out.splitlines()[:80] if out else []
    nt = os.path.expanduser("~/nuclei-templates")
    if os.path.exists(nt) and shutil_which("nuclei"):
        rc, out = run_cmd(f"nuclei -u {target} -t {nt} -severity high -o {os.path.join(outdir,'nuclei_high.txt')}", timeout=1200)
        ag['nuclei_high'] = "saved"
    if shutil_which("sqlmap"):
        rc, out = run_cmd(f"sqlmap -u \"{target}\" --batch --crawl=1 --level=1 --risk=1", timeout=1800)
        try:
            with open(os.path.join(outdir, "sqlmap_quick.txt"), "w", encoding="utf-8") as fh:
                fh.write(out or "")
        except Exception:
            pass
        ag['sqlmap'] = "saved"
    return ag

def create_word_report(target, outdir, summary):
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    doc.add_heading(f'OWASP Top10 Scan Report - {target}', level=1)
    doc.add_paragraph(f"Generated: {datetime.datetime.utcnow().isoformat()} UTC")
    doc.add_paragraph("Note: Passive/Non-destructive checks by default. Active checks only if --active used with permission.\n")
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(summary.get('exec', 'No executive summary provided.'))
    doc.add_heading("Findings by OWASP Category", level=2)
    for cat, findings in summary.get('owasp', {}).items():
        doc.add_heading(cat, level=3)
        if not findings:
            doc.add_paragraph("No obvious findings (in passive checks).")
            continue
        for f in findings:
            doc.add_paragraph(json.dumps(f, ensure_ascii=False)[:1000])
    doc.add_heading("Raw Tool Outputs (selected)", level=2)
    for fname in sorted(os.listdir(outdir)):
        if fname.endswith(".txt") or fname.endswith(".json"):
            p = os.path.join(outdir, fname)
            doc.add_heading(fname, level=4)
            try:
                with open(p, "r", encoding="utf-8", errors="ignore") as fh:
                    txt = fh.read(4000)
                    doc.add_paragraph(txt)
            except Exception:
                doc.add_paragraph(f"Could not read {fname}")
    doc_path = os.path.join(outdir, f"OWASP_Scan_Report_{datetime.datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}.docx")
    doc.save(doc_path)
    return doc_path

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 owasp_scan_report.py <target_url> [--active]")
        sys.exit(1)
    target = sys.argv[1].rstrip('/')
    active = "--active" in sys.argv
    parsed = urlparse(target)
    if not parsed.scheme:
        target = "http://" + target
        parsed = urlparse(target)
    timestamp = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    outdir = os.path.join(RESULTS_DIR, f"{parsed.netloc.replace(':','_')}_{timestamp}")
    ensure_dir(outdir)
    summary = {'target': target, 'timestamp': timestamp, 'exec': '', 'owasp': {}}
    summary['basic'] = basic_tech_info(target, outdir)
    hfind = check_headers_tls(target, outdir)
    summary['basic']['tls_headers'] = hfind
    links, forms = find_links(target, outdir)
    summary['basic']['links_count'] = len(links)
    summary['sensitive_files'] = check_sensitive_files(target, outdir)
    summary['owasp']['A1-Injection'] = check_sqli_basic(target, links, outdir)
    summary['owasp']['A2-BrokenAuth'] = []
    headers = hfind.get('headers', {}) if isinstance(hfind, dict) else {}
    try:
        summary['owasp']['A2-BrokenAuth'] = []
        if headers:
            sc = headers.get('set-cookie', '')
            if sc:
                summary['owasp']['A2-BrokenAuth'].append({'set-cookie': sc})
    except Exception:
        pass
    summary['owasp']['A3-SensitiveDataExposure'] = []
    if 'tls_verify' in hfind:
        summary['owasp']['A3-SensitiveDataExposure'].append({'tls': hfind.get('tls_verify')})
    summary['owasp']['A4-XXE'] = []
    for l in links:
        if l.lower().endswith('.xml') or 'xml' in l.lower() or 'wsdl' in l.lower():
            summary['owasp']['A4-XXE'].append({'endpoint': l})
    summary['owasp']['A5-BrokenAccessControl'] = check_idor(target, links, outdir)
    summary['owasp']['A6-SecurityMisconfig'] = check_security_misconfig(hfind.get('headers', {}) if isinstance(hfind, dict) else {}, outdir)
    summary['owasp']['A7-XSS'] = check_xss_reflected(target, links, forms, outdir)
    summary['owasp']['A8-InsecureDeserial'] = []
    summary['owasp']['A9-ComponentsWithVuln'] = summary['basic'].get('whatweb', []) + summary['basic'].get('nmap', [])
    summary['owasp']['A10-LoggingAndMonitoring'] = []
    summary['csrf'] = check_csrf(forms, outdir)
    if active:
        print("[!] ACTIVE MODE ENABLED. Make sure you have explicit written permission to run aggressive tests.")
        ag = aggressive_checks(target, outdir)
        summary['aggressive'] = ag
    execsum = []
    for k, v in summary['owasp'].items():
        if v:
            execsum.append(f"{k}: {len(v)} finding(s)")
    summary['exec'] = " | ".join(execsum) if execsum else "No obvious findings in passive checks."
    try:
        with open(os.path.join(outdir, "summary.json"), "w", encoding="utf-8") as fh:
            fh.write(json.dumps(summary, indent=2, ensure_ascii=False))
    except Exception:
        pass
    docx = create_word_report(target, outdir, summary)
    print("[+] Report created:", docx)
    print("[+] All outputs in:", outdir)

if __name__ == "__main__":
    main()
